import io
import pyttsx3
from PyPDF2 import PdfFileReader

# metni sesli olarak okuyan fonksiyon
def read_text(text):
    engine = pyttsx3.init()
    # Türkçe seslendirme için kullanılacak sesi seçiyoruz
    voices = engine.getProperty('voices')
    for voice in voices:
        if voice.name == "Turkish":
            engine.setProperty('voice', voice.id)
            break
    # metni okuyoruz
    engine.say(text)
    engine.runAndWait()

# pdf dosyasını okuyan fonksiyon
def read_pdf(filename):
    with open(filename, "rb") as f:
        pdf = PdfFileReader(f)
        total_pages = pdf.getNumPages()
        text = ""
        for i in range(total_pages):
            page = pdf.getPage(i)
            text += page.extractText()
        read_text(text)

# txt dosyasını okuyan fonksiyon
def read_txt(filename):
    with io.open(filename, "r", encoding="utf-8") as f:
        text = f.read()
        read_text(text)

# docx dosyasını okuyan fonksiyon
def read_docx(filename):
    try:
        import docx
    except ImportError:
        print("python-docx kütüphanesi yüklü değil. Lütfen yükleyin!")
        return
    doc = docx.Document(filename)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    read_text(text)

# odt dosyasını okuyan fonksiyon
def read_odt(filename):
    try:
        import odf
    except ImportError:
        print("odfpy kütüphanesi yüklü değil. Lütfen yükleyin!")
        return
    from odf import text
    with io.open(filename, "rb") as f:
        doc = text.load(f)
        text = "\n".join([element.text for element in doc.getElementsByType(text.P)])
        read_text(text)

# rtf dosyasını okuyan fonksiyon
def read_rtf(filename):
    try:
        import pyth.plugins.plaintext.writer
    except ImportError:
        print("pyth kütüphanesi yüklü değil. Lütfen yükleyin!")
        return
    import pyth
    with io.open(filename, "rb") as f:
        doc = pyth.Document()
        doc.load_rtf(f)
        text = pyth.plugins.plaintext.writer.PlaintextWriter.write(doc).getvalue()
        read_text(text)

# kullanıcıdan dosya adını al
file_name = input("Lütfen okunacak dosyanın adını giriniz: ")

# dosya uzantısına göre okuma işlemini gerçekleştir
if file_name.endswith(".pdf"):
    read_pdf(file_name)
elif file_name.endswith(".txt"):
    read_txt(file_name)
elif file_name.endswith(".docx"):
    read_docx(file_name)
elif file_name.endswith(".odt"):
    read_odt(file_name)
elif file_name.endswith(".rtf"):
    read_rtf(file_name)
else:
    print("Desteklenmeyen dosya türü!")
